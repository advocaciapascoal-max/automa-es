"""
Motor de formatacao FIEL ao padrao de pecas do escritorio.

Especificacoes (padrao vigente desde 31/08/2026 - medido sobre o acervo real):
  Pagina A4
  Margens: topo 3,5 cm | base 3,0 cm | esquerda 3,0 cm | direita 3,0 cm
  Cabecalho e rodape: 1,25 cm
  Fonte: Verdana | Corpo 11 pt | Citacoes 10 pt italico
  Espacamento entre linhas: 1,5
  Recuo 1a linha do corpo: 3,0 cm
  Recuo esquerdo das citacoes: 4 cm
  Alinhamento padrao: justificado
  Enderecamento: bold, justificado, sem recuo
  Nome da peca: centralizado, bold
  Titulos de secao: CAIXA ALTA, CENTRALIZADOS, bold, SEM numeracao e SEM letra
    (nunca "A -", nunca "I -"; numeracao legada e removida automaticamente)
  Local/data: alinhado a direita
  Assinaturas (nome + OAB): centralizadas, bold
  Negrito de corpo: so a oracao-chave do paragrafo, nunca o paragrafo inteiro

ATENCAO: o acervo antigo em "Z:/Advocacia Pascoal/1 DOCS" esta com 1,15 e 2,5 cm.
Em 31/08/2026 o Dr. Alexandre corrigiu expressamente: e 3 cm e 1,5. Nao "corrigir"
de volta para o que esta medido nas pecas antigas.

Mantem o timbrado (cabecalho/logo/rodape) intacto - apenas limpa o body.
O timbrado e fornecido pelo escritorio em config/timbrado_modelo.docx.
"""
import logging
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

_log = logging.getLogger('agente_op.formatacao')

# Timbrado fornecido pelo escritorio
_ROOT = Path(__file__).resolve().parent.parent.parent
TEMPLATE_PATH = _ROOT / 'config' / 'timbrado_modelo.docx'
REFERENCIAS_DIR = Path(__file__).parent / 'REFERENCIAS'

FONTE = 'Verdana'
TAM_CORPO = Pt(11)
TAM_CITACAO = Pt(10)
RECUO_PRIMEIRA = Cm(3.0)
RECUO_CITACAO_ESQ = Cm(4.0)
LINE_SPACING = 1.5
MARG_TOP = Cm(3.5)
MARG_BOT = Cm(3.0)
MARG_LEFT = Cm(3.0)
MARG_RIGHT = Cm(3.0)
DIST_CABECALHO = Cm(1.25)
DIST_RODAPE = Cm(1.25)

# Respiro entre paragrafos. Nao consta do padrao fixado pelo escritorio;
# zerar aqui caso se prefira paragrafos colados.
ESPACO_APOS = Pt(6)
ESPACO_TITULO = Pt(12)


# ============================================================
# DETECTORES DE TIPO DE BLOCO
# ============================================================

NOMES_PECA_CENTRALIZADO = {
    'RECLAMACAO TRABALHISTA', 'RECLAMATORIA TRABALHISTA',
    'CONTESTACAO',
    'REPLICA',
    'RAZOES FINAIS',
    'MEMORIAIS', 'MEMORIAIS ESCRITOS',
    'RECURSO ORDINARIO', 'RAZOES DO RECURSO ORDINARIO',
    'RAZOES DE RECURSO ORDINARIO',
    'RECURSO DE REVISTA', 'RAZOES DO RECURSO DE REVISTA',
    'CONTRARRAZOES', 'CONTRARRAZOES AO RECURSO ORDINARIO',
    'CONTRARRAZOES AO RECURSO DE REVISTA',
    'EMBARGOS DE DECLARACAO',
    'AGRAVO DE INSTRUMENTO', 'AGRAVO INTERNO', 'AGRAVO DE PETICAO',
    'PETICAO INICIAL',
    'MANIFESTACAO',
    'PARECER', 'PARECER JURIDICO',
}

# Numeracao legada de titulo que deve ser removida: "I - ", "IV.I - ", "01 – ", "A- "
_RE_NUMERACAO_TITULO = re.compile(
    r'^(?:[IVXLCDM]+(?:\.[IVXLCDM\d]+)*|\d+(?:\.\d+)*|[A-Z])\s*[\.\-–—)]\s*',
)

_RE_LOCAL_DATA = re.compile(
    r'^(?:Sorocaba|Campinas|S[ãa]o Paulo|Jacare[íi]|Bras[íi]lia)\s*[,/]\s*\S',
)
_RE_OAB = re.compile(r'^OAB[/\s]', re.IGNORECASE)
_RE_CABECALHO_PROCESSUAL = re.compile(
    r'^(?:Origem|Processo|PROCESSO|Recorrente|Recorrida|Recorridas|Recorrido|Recorridos'
    r'|Reclamante|Reclamada|Reclamadas|Embargante|Embargada|Autor|Autora|R[ée]|R[ée]u)'
    r'\s*(?:n[º°ºo]|N[º°O]|:)',
)
_RE_VOCATIVO = re.compile(
    r'^(?:Egr[ée]gio Tribunal|Colendo Tribunal|Doutos Julgadores|Eminentes Julgadores'
    r'|Colenda Turma|Excelent[íi]ssimos Senhores Desembargadores)\s*[,\.]?$',
)

FECHO_CENTRALIZADO = {
    'nestes termos,', 'nestes termos', 'termos em que,', 'termos em que',
    'pede deferimento.', 'pede deferimento', 'p. deferimento.', 'p. deferimento',
    'pede e espera deferimento.', 'pede e aguarda deferimento.',
}

_RE_CITACAO_INICIO = re.compile(
    r'^(?:["“«]|s[uú]mula\s+\d+|oj\s+\d+|tese\s+\d+|tema\s+\d+'
    r'|art(?:igo)?\.?\s+\d+(?:[,\.\-\s][^—\-:]*)?[—\-:])',
    re.IGNORECASE,
)

_RE_PEDIDO_FINAL = re.compile(r'^[a-z]{1,2}\)\s+')


def _normalizar_acento(s):
    tabela = str.maketrans('ÇÁÀÃÂÉÊÍÓÕÔÚÜ', 'CAAAAEEIOOOUU')
    return s.upper().translate(tabela)


def _eh_caixa_alta(s):
    letras = [c for c in s if c.isalpha()]
    return bool(letras) and all(c.isupper() for c in letras)


def _eh_enderecamento(linha):
    s = linha.strip().upper()
    return s.startswith(('EXCELENTISSIMO', 'EXCELENTÍSSIMO', 'EXMO', 'EXMA',
                         'AO EGREGIO', 'AO EGRÉGIO'))


def _eh_cabecalho_processual(linha):
    return bool(_RE_CABECALHO_PROCESSUAL.match(linha.strip()))


def _eh_nome_peca(linha):
    s = re.sub(r'[\.\:\-—\s]+$', '', linha.strip())
    return _normalizar_acento(s) in NOMES_PECA_CENTRALIZADO


def _limpar_numeracao_titulo(texto):
    """Remove numeracao/letra legada do titulo ('I - DA X' -> 'DA X')."""
    limpo = _RE_NUMERACAO_TITULO.sub('', texto.strip(), count=1)
    return limpo if len(limpo) >= 3 else texto.strip()


def _eh_titulo_secao(linha):
    """Titulo e linha curta em CAIXA ALTA, com ou sem numeracao legada."""
    s = linha.strip()
    if not (6 <= len(s) <= 200):
        return False
    if s.endswith(('.', ':', ';', ',')):
        return False
    return _eh_caixa_alta(_limpar_numeracao_titulo(s))


def _eh_citacao(linha):
    s = linha.strip()
    if not s:
        return False
    if s.startswith(('"', '“', '«')):
        return True
    return len(s) >= 30 and bool(_RE_CITACAO_INICIO.match(s))


def _eh_local_data(linha):
    return bool(_RE_LOCAL_DATA.match(linha.strip()))


def _eh_bloco_assinatura(linhas):
    """Bloco de assinatura: nome numa linha, OAB na seguinte."""
    return len(linhas) == 2 and bool(_RE_OAB.match(linhas[1].strip()))


def _eh_fecho(linhas):
    return all(l.strip().lower() in FECHO_CENTRALIZADO for l in linhas)


def _eh_vocativo(linha):
    return bool(_RE_VOCATIVO.match(linha.strip()))


def _eh_pedido_final(linha):
    return bool(_RE_PEDIDO_FINAL.match(linha.strip()))


# ============================================================
# RENDERIZADORES
# ============================================================

def _ajustar_secoes(doc):
    for s in doc.sections:
        s.top_margin = MARG_TOP
        s.bottom_margin = MARG_BOT
        s.left_margin = MARG_LEFT
        s.right_margin = MARG_RIGHT
        s.header_distance = DIST_CABECALHO
        s.footer_distance = DIST_RODAPE


def _limpar_corpo(doc):
    body = doc.element.body
    for el in list(body):
        if el.tag.split('}')[-1] in ('p', 'tbl'):
            body.remove(el)


def _estilo_run(run, *, size=TAM_CORPO, bold=False, italic=False):
    run.font.name = FONTE
    run.font.size = size
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = italic


def _padrao_paragrafo(p, *, alinhamento=WD_ALIGN_PARAGRAPH.JUSTIFY,
                      recuo_primeira=False, left_indent=None,
                      espaco_antes=Pt(0), espaco_apos=ESPACO_APOS):
    p.alignment = alinhamento
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_before = espaco_antes
    pf.space_after = espaco_apos
    pf.first_line_indent = RECUO_PRIMEIRA if recuo_primeira else Cm(0)
    pf.left_indent = left_indent if left_indent is not None else Cm(0)


def _add_texto_inline(p, texto, *, size=TAM_CORPO, negrito_total=False, italic_total=False):
    """Adiciona texto a um paragrafo respeitando markup **bold** e _italic_."""
    pos = 0
    pattern = re.compile(r'(\*\*([^*]+)\*\*|_([^_\n]+)_)')
    for m in pattern.finditer(texto):
        if m.start() > pos:
            r = p.add_run(texto[pos:m.start()])
            _estilo_run(r, size=size, bold=negrito_total, italic=italic_total)
        if m.group(2):
            r = p.add_run(m.group(2))
            _estilo_run(r, size=size, bold=True, italic=italic_total)
        elif m.group(3):
            r = p.add_run(m.group(3))
            _estilo_run(r, size=size, bold=negrito_total, italic=True)
        pos = m.end()
    if pos < len(texto):
        r = p.add_run(texto[pos:])
        _estilo_run(r, size=size, bold=negrito_total, italic=italic_total)


def _add_paragrafo_corpo(doc, texto):
    p = doc.add_paragraph()
    _padrao_paragrafo(p, recuo_primeira=True)
    _add_texto_inline(p, texto)
    return p


def _add_paragrafo_titulo_secao(doc, texto):
    """Titulo de secao: CAIXA ALTA, centralizado, bold, sem numeracao."""
    p = doc.add_paragraph()
    _padrao_paragrafo(p, alinhamento=WD_ALIGN_PARAGRAPH.CENTER,
                      espaco_antes=ESPACO_TITULO, espaco_apos=ESPACO_TITULO)
    _add_texto_inline(p, _limpar_numeracao_titulo(texto).upper(), negrito_total=True)
    return p


def _add_paragrafo_nome_peca(doc, texto):
    p = doc.add_paragraph()
    _padrao_paragrafo(p, alinhamento=WD_ALIGN_PARAGRAPH.CENTER,
                      espaco_antes=ESPACO_TITULO, espaco_apos=ESPACO_TITULO)
    _add_texto_inline(p, texto.upper(), negrito_total=True)
    return p


def _add_paragrafo_enderecamento(doc, texto):
    p = doc.add_paragraph()
    _padrao_paragrafo(p)
    _add_texto_inline(p, texto, negrito_total=True)
    return p


def _add_paragrafo_cabecalho(doc, linhas):
    """Origem / Processo / Recorrente / Recorrido: a esquerda, sem recuo."""
    negrito = linhas[0].strip().upper().startswith('PROCESSO')
    for linha in linhas:
        p = doc.add_paragraph()
        _padrao_paragrafo(p, alinhamento=WD_ALIGN_PARAGRAPH.LEFT, espaco_apos=Pt(0))
        _add_texto_inline(p, linha.strip(), negrito_total=negrito)


def _add_paragrafo_vocativo(doc, linhas):
    for linha in linhas:
        p = doc.add_paragraph()
        _padrao_paragrafo(p, alinhamento=WD_ALIGN_PARAGRAPH.LEFT, espaco_apos=Pt(0))
        _add_texto_inline(p, linha.strip())


def _add_paragrafo_pedido(doc, texto):
    """Pedidos finais a), b), c): justificado, recuo de 1a linha."""
    p = doc.add_paragraph()
    _padrao_paragrafo(p, recuo_primeira=True)
    _add_texto_inline(p, texto)
    return p


def _add_paragrafo_citacao(doc, texto):
    """Citacao em bloco: justificado, recuo esquerdo 4cm, 10pt, italico."""
    p = doc.add_paragraph()
    _padrao_paragrafo(p, left_indent=RECUO_CITACAO_ESQ)
    _add_texto_inline(p, texto, size=TAM_CITACAO, italic_total=True)
    return p


def _add_paragrafo_local_data(doc, texto):
    p = doc.add_paragraph()
    _padrao_paragrafo(p, alinhamento=WD_ALIGN_PARAGRAPH.RIGHT, espaco_antes=Pt(18))
    _add_texto_inline(p, texto)
    return p


def _add_bloco_assinatura(doc, linhas):
    """Nome + OAB: centralizados, bold."""
    for linha in linhas:
        p = doc.add_paragraph()
        _padrao_paragrafo(p, alinhamento=WD_ALIGN_PARAGRAPH.CENTER, espaco_apos=Pt(0))
        _add_texto_inline(p, linha.strip(), negrito_total=True)
    vazio = doc.add_paragraph()
    _padrao_paragrafo(vazio, espaco_apos=Pt(0))


def _add_bloco_fecho(doc, linhas):
    for linha in linhas:
        p = doc.add_paragraph()
        _padrao_paragrafo(p, alinhamento=WD_ALIGN_PARAGRAPH.CENTER, espaco_apos=Pt(0))
        _add_texto_inline(p, linha.strip())


def adicionar_imagens_anexo(doc, imagens, drive_service):
    """
    Acrescenta secao de ANEXOS ao final do documento com imagens do Drive.

    imagens: lista de {'id', 'name', 'mime'} | drive_service: client Drive autenticado.
    Falha de uma imagem nao interrompe as demais - apenas registra no log.
    """
    import io as _io
    from docx.enum.text import WD_BREAK
    from googleapiclient.http import MediaIoBaseDownload

    if not imagens:
        return

    quebra = doc.add_paragraph()
    quebra.add_run().add_break(WD_BREAK.PAGE)

    titulo = doc.add_paragraph()
    _padrao_paragrafo(titulo, alinhamento=WD_ALIGN_PARAGRAPH.CENTER,
                      espaco_antes=ESPACO_TITULO, espaco_apos=ESPACO_TITULO)
    _add_texto_inline(titulo, 'ANEXOS – PROVAS DOCUMENTAIS (IMAGENS)', negrito_total=True)

    for i, img in enumerate(imagens, 1):
        try:
            buf = _io.BytesIO()
            req = drive_service.files().get_media(fileId=img['id'], supportsAllDrives=True)
            dl = MediaIoBaseDownload(buf, req)
            concluido = False
            while not concluido:
                _, concluido = dl.next_chunk()
            buf.seek(0)

            legenda = doc.add_paragraph()
            _padrao_paragrafo(legenda, alinhamento=WD_ALIGN_PARAGRAPH.CENTER,
                              espaco_apos=Pt(4))
            _add_texto_inline(legenda, f'**ANEXO {i}** — {img["name"]}')

            p_img = doc.add_paragraph()
            _padrao_paragrafo(p_img, alinhamento=WD_ALIGN_PARAGRAPH.CENTER)
            try:
                p_img.add_run().add_picture(buf, width=Cm(15))
            except Exception as e:
                _log.warning('falha ao inserir imagem %s: %s', img['name'], e)
                continue
            _log.info('  anexo %d inserido: %s', i, img['name'])
        except Exception as e:
            _log.warning('falha ao baixar/inserir imagem %s: %s', img['name'], e)


# ============================================================
# RENDERIZADOR PRINCIPAL
# ============================================================

def montar_documento(texto_peca: str):
    """Monta o Document no timbrado, formatado, SEM salvar.

    Util para quem precisa acrescentar conteudo (anexos, por exemplo) antes de gravar.
    """
    if not Path(TEMPLATE_PATH).exists():
        raise FileNotFoundError(
            f'Timbrado nao encontrado: {TEMPLATE_PATH}. '
            f'O escritorio deve fornecer o timbrado DELE em config/timbrado_modelo.docx.'
        )

    doc = Document(str(TEMPLATE_PATH))
    _ajustar_secoes(doc)
    _limpar_corpo(doc)

    # Quebra em blocos (separados por linhas em branco)
    blocos, atual = [], []
    for linha in texto_peca.split('\n'):
        if linha.strip():
            atual.append(linha.rstrip())
        elif atual:
            blocos.append(atual)
            atual = []
    if atual:
        blocos.append(atual)

    for bloco in blocos:
        unica = len(bloco) == 1
        primeira = bloco[0].strip()
        corrido = ' '.join(l.strip() for l in bloco)

        if _eh_enderecamento(primeira):
            _add_paragrafo_enderecamento(doc, corrido)
        elif all(_eh_cabecalho_processual(l) for l in bloco):
            _add_paragrafo_cabecalho(doc, bloco)
        elif unica and _eh_nome_peca(primeira):
            _add_paragrafo_nome_peca(doc, primeira)
        elif unica and _eh_titulo_secao(primeira):
            _add_paragrafo_titulo_secao(doc, primeira)
        elif _eh_citacao(primeira):
            _add_paragrafo_citacao(doc, corrido)
        elif unica and _eh_local_data(primeira):
            _add_paragrafo_local_data(doc, primeira)
        elif _eh_bloco_assinatura(bloco):
            _add_bloco_assinatura(doc, bloco)
        elif _eh_fecho(bloco):
            _add_bloco_fecho(doc, bloco)
        elif all(_eh_vocativo(l) for l in bloco):
            _add_paragrafo_vocativo(doc, bloco)
        elif unica and _eh_pedido_final(primeira):
            _add_paragrafo_pedido(doc, primeira)
        else:
            _add_paragrafo_corpo(doc, corrido)

    return doc


def gerar_peca_escritorio(texto_peca: str, output_path, imagens=None, drive_service=None):
    """Gera DOCX no timbrado seguindo a formatacao padrao do escritorio.

    imagens/drive_service (opcionais): anexa imagens da pasta do cliente ao final,
    em secao propria depois de quebra de pagina.
    """
    doc = montar_documento(texto_peca)
    if imagens:
        adicionar_imagens_anexo(doc, imagens, drive_service)
    out = Path(output_path)
    doc.save(str(out))
    return out
